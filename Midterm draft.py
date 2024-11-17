import hashlib
import pandas as pd
from datetime import datetime
from datetime import datetime, timedelta
from docx import Document
from docx.shared import Pt  # For setting font sizes
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT  # For paragraph alignment
from fpdf import FPDF
import openai
import matplotlib.pyplot as plt


# Path to the Excel file
file_path = r'C:\Users\bubad\Desktop\attendance app data.xlsx'


# Hardcoded user data with encrypted passwords
users = {
    "user": hashlib.sha256("password".encode()).hexdigest(),
    "user2": hashlib.sha256("password2".encode()).hexdigest()
}

# Sample data structures for courses and students
courses = {}
students = {}

# Function to encrypt password
def encrypt_password(password):
    return hashlib.sha256(password.encode()).hexdigest()

# Function to log in
def login():
    while True:
        print("===== Login Menu =====")
        print("1. Log in")
        print("2. Exit")
        choice = input("Enter your choice: ")

        if choice == "1":
            username = input("Enter username: ")
            password = input("Enter password: ")
            encrypted_password = encrypt_password(password)

            # Validate credentials
            if username in users and users[username] == encrypted_password:
                print(f"Welcome, {username}!")
                logged_in_menu(username)
            else:
                print("Invalid username or password. Please try again.")
        elif choice == "2":
            print("Exiting the application...")
            break
        else:
            print("Invalid choice. Please select again.")

# Function to take attendance and automatically mark students as absent if they don't check in
def take_attendance():
    print("===== Take Attendance =====")
    
    # Load the 'Courses' sheet to dynamically pull the list of courses
    courses_df = pd.read_excel(file_path, sheet_name='Courses')
    course_names = courses_df['course_name'].tolist()

    # Ask the user to select a course from the list
    print("Select a course:")
    for index, course_name in enumerate(course_names, 1):
        print(f"{index}. {course_name}")
    course_choice = int(input("Enter the number of the course: ")) - 1

    if course_choice < 0 or course_choice >= len(course_names):
        print("Invalid course selection.")
        return

    course_name = course_names[course_choice]
    print(f"Recording attendance for {course_name}")

    # Load the 'All students' sheet
    all_students_df = pd.read_excel(file_path, sheet_name='All students')

    # Filter students enrolled in the selected course
    enrolled_students = all_students_df[all_students_df['course_name'] == course_name]

    # Track students who checked in
    checked_in_students = set()

    # Load the existing attendance sheet or create an empty one if it doesn't exist
    try:
        attendance_df = pd.read_excel(file_path, sheet_name='Attendance')
    except FileNotFoundError:
        attendance_df = pd.DataFrame(columns=['student_id', 'first_name', 'last_name', 'course_name', 'check_in_datetime', 'status'])

    # Get the course start time from the 'Courses' sheet and extract only the start time part
    meeting_time = courses_df.loc[courses_df['course_name'] == course_name, 'meeting_time'].values[0]
    course_start_time_str = meeting_time.split('-')[0].strip()  # Split on dash and strip spaces around it
    
    # Debugging print to check the extracted start time string
    print(f"Extracted course start time string: '{course_start_time_str}'")

    # Convert start time to datetime object
    try:
        course_start_time = datetime.strptime(course_start_time_str, '%I:%M %p')
    except ValueError as e:
        print(f"Error converting course start time: {e}")
        return

    while True:
        # Ask for student ID
        student_id = input("Enter student ID to check in (or type 'Exit' to finish attendance): ")

        if student_id.lower() == 'exit':
            print("Finishing attendance and marking absentees.")
            break

        # Ensure student_id is being treated as a string and handle lookup
        student_row = enrolled_students[enrolled_students['student_id'].astype(str) == student_id]

        if not student_row.empty:
            # Extract the first and last name of the student
            first_name = student_row.iloc[0]['first_name']
            last_name = student_row.iloc[0]['last_name']

            # Log the current time for check-in
            check_in_datetime = datetime.now()
            check_in_time_str = check_in_datetime.strftime("%Y-%m-%d %H:%M:%S")

            # Determine attendance status (late if more than 10 minutes after the class start time)
            if check_in_datetime.time() > (course_start_time + timedelta(minutes=10)).time():
                status = 'Late'
            else:
                status = 'Present'

            # Mark the student as checked in
            checked_in_students.add(student_id)

            # Log the attendance record
            new_attendance = pd.DataFrame({
                'student_id': [student_id],
                'first_name': [first_name],
                'last_name': [last_name],
                'course_name': [course_name],
                'check_in_datetime': [check_in_time_str],
                'status': [status]
            })

            # Append the new attendance record to the DataFrame
            attendance_df = pd.concat([attendance_df, new_attendance], ignore_index=True)

            print(f"Attendance captured for {first_name} {last_name} (ID: {student_id}) as {status} at {check_in_time_str}.")
        else:
            print("Student ID not found. Please try again.")

    # Automatically mark students as 'Absent' if they haven't checked in
    absent_students = enrolled_students[~enrolled_students['student_id'].astype(str).isin(checked_in_students)]

    for index, absent_student in absent_students.iterrows():
        absent_attendance = pd.DataFrame({
            'student_id': [absent_student['student_id']],
            'first_name': [absent_student['first_name']],
            'last_name': [absent_student['last_name']],
            'course_name': [course_name],
            'check_in_datetime': [datetime.now().strftime("%Y-%m-%d %H:%M:%S")],
            'status': ['Absent']
        })

        attendance_df = pd.concat([attendance_df, absent_attendance], ignore_index=True)

        print(f"Marked {absent_student['first_name']} {absent_student['last_name']} (ID: {absent_student['student_id']}) as Absent.")

    # Save the updated attendance DataFrame back to the Excel file
    with pd.ExcelWriter(file_path, mode='a', engine='openpyxl', if_sheet_exists='replace') as writer:
        attendance_df.to_excel(writer, sheet_name='Attendance', index=False)

    print(f"Attendance for {course_name} has been recorded.")


# Function to add a course
def add_course():
    print("===== Add New Course =====")
    
    # Prompt the user to input the course details, with the option to cancel
    course_name = input("Enter course name (or type 'cancel' to exit): ")
    if course_name.lower() == 'cancel':
        print("Operation canceled.")
        return

    meeting_time = input("Enter meeting time (e.g., 10:00 AM - 11:30 AM) (or type 'cancel' to exit): ")
    if meeting_time.lower() == 'cancel':
        print("Operation canceled.")
        return

    meeting_date = input("Enter meeting dates (e.g., Mondays and Wednesdays) (or type 'cancel' to exit): ")
    if meeting_date.lower() == 'cancel':
        print("Operation canceled.")
        return

    professor = input("Enter professor's name (or type 'cancel' to exit): ")
    if professor.lower() == 'cancel':
        print("Operation canceled.")
        return

    description = input("Enter course description (or type 'cancel' to exit): ")
    if description.lower() == 'cancel':
        print("Operation canceled.")
        return

    # Load the 'Courses' sheet
    courses_df = pd.read_excel(file_path, sheet_name='Courses')

    # Create a new row for the course with the required columns
    new_course = pd.DataFrame({
        'course_name': [course_name],
        'meeting_time': [meeting_time],
        'meeting_date': [meeting_date],
        'professor': [professor],
        'description': [description]
    })

    # Append the new course to the DataFrame
    courses_df = pd.concat([courses_df, new_course], ignore_index=True)

    # Save the updated DataFrame back to the Excel file
    with pd.ExcelWriter(file_path, mode='a', engine='openpyxl', if_sheet_exists='replace') as writer:
        courses_df.to_excel(writer, sheet_name='Courses', index=False)

    print(f"Course {course_name} has been added successfully.")

def add_student():
    print("===== Add Student =====")

    # Load the 'All students' and 'Courses' sheets
    all_students_df = pd.read_excel(file_path, sheet_name='All students')
    courses_df = pd.read_excel(file_path, sheet_name='Courses')

    # Get the list of majors and courses
    majors = all_students_df['major'].unique().tolist()
    courses = courses_df['course_name'].tolist()

    # Ask for student details with a cancel option
    student_id = input("Enter student ID (or type 'cancel' to exit): ")
    if student_id.lower() == 'cancel':
        print("Operation canceled.")
        return

    first_name = input("Enter first name (or type 'cancel' to exit): ")
    if first_name.lower() == 'cancel':
        print("Operation canceled.")
        return

    last_name = input("Enter last name (or type 'cancel' to exit): ")
    if last_name.lower() == 'cancel':
        print("Operation canceled.")
        return

    email = input("Enter email (or type 'cancel' to exit): ")
    if email.lower() == 'cancel':
        print("Operation canceled.")
        return

    # Let the user select a major
    print("\nSelect a major from the list (or type 'cancel' to exit):")
    for idx, major in enumerate(majors, 1):
        print(f"{idx}. {major}")

    try:
        major_choice = input("Enter the number of the major: ")
        if major_choice.lower() == 'cancel':
            print("Operation canceled.")
            return
        major_choice = int(major_choice) - 1
        if major_choice < 0 or major_choice >= len(majors):
            print("Invalid major selection.")
            return
    except ValueError:
        print("Invalid input. Please enter a valid number.")
        return
    selected_major = majors[major_choice]

    # Let the user select a course
    print("\nSelect a course from the list (or type 'cancel' to exit):")
    for idx, course in enumerate(courses, 1):
        print(f"{idx}. {course}")

    try:
        course_choice = input("Enter the number of the course: ")
        if course_choice.lower() == 'cancel':
            print("Operation canceled.")
            return
        course_choice = int(course_choice) - 1
        if course_choice < 0 or course_choice >= len(courses):
            print("Invalid course selection.")
            return
    except ValueError:
        print("Invalid input. Please enter a valid number.")
        return
    selected_course = courses[course_choice]

    # Append the new student to the DataFrame
    new_student = pd.DataFrame({
        'student_id': [student_id],
        'first_name': [first_name],
        'last_name': [last_name],
        'email': [email],
        'major': [selected_major],
        'course_name': [selected_course],
        'is_deleted': [False],
        'deletion_date': [None]
    })

    # Concatenate the new student data to the existing student DataFrame
    all_students_df = pd.concat([all_students_df, new_student], ignore_index=True)

    # Save the updated DataFrame back to the Excel file
    with pd.ExcelWriter(file_path, mode='a', engine='openpyxl', if_sheet_exists='replace') as writer:
        all_students_df.to_excel(writer, sheet_name='All students', index=False)

    print(f"\nStudent {first_name} {last_name} has been added successfully to {selected_course}.")




# Function to generate a report by student this ask if the user would like to save the doc as well
def report_by_student():
    student_id = input("Enter the student ID to generate the report: ")

    # Load the 'Attendance' and 'All students' sheets
    attendance_df = pd.read_excel(file_path, sheet_name='Attendance')
    all_students_df = pd.read_excel(file_path, sheet_name='All students')

    # Check if the student exists
    student_row = all_students_df[all_students_df['student_id'].astype(str) == student_id]
    if student_row.empty:
        print(f"No student found with ID {student_id}.")
        return

    # Get student information
    first_name = student_row.iloc[0]['first_name']
    last_name = student_row.iloc[0]['last_name']

    # Filter attendance records for the student
    student_attendance = attendance_df[attendance_df['student_id'].astype(str) == student_id]

    if student_attendance.empty:
        print(f"No attendance records found for student {first_name} {last_name} (ID: {student_id}).")
        return

    # Fill 'nan' status values with 'Unknown'
    student_attendance['status'].fillna('Unknown', inplace=True)

    # Count attendance statistics for the student
    present_count = student_attendance[student_attendance['status'] == 'Present'].shape[0]
    absent_count = student_attendance[student_attendance['status'] == 'Absent'].shape[0]
    late_count = student_attendance[student_attendance['status'] == 'Late'].shape[0]

    # Display the report on the console
    print(f"Attendance Report for {first_name} {last_name} (ID: {student_id})")
    print(f"Present: {present_count}")
    print(f"Absent: {absent_count}")
    print(f"Late: {late_count}")

    print("\nDetailed Attendance Records:")
    for index, row in student_attendance.iterrows():
        date = row['check_in_datetime']
        status = row['status']
        course_name = row['course_name']
        print(f"{date} - {course_name}: {status}")

    # Ask if the user wants to export the report to a Word document
    export_choice = input("\nWould you like to export this report to a Word document? (y/n): ").lower()

    if export_choice == 'y':
        # Generate Word Document
        doc = Document()

        # Add header
        header = doc.sections[0].header
        header_paragraph = header.paragraphs[0]
        header_paragraph.text = f"Attendance Report - {first_name} {last_name}"
        header_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Add title
        doc.add_heading(f"Attendance Report for {first_name} {last_name} (ID: {student_id})", level=1)

        # Add report date
        doc.add_paragraph(f"Report generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")

        # Add attendance stats
        doc.add_paragraph("Attendance Summary:")
        doc.add_paragraph(f"Present: {present_count}").bold = True
        doc.add_paragraph(f"Absent: {absent_count}").bold = True
        doc.add_paragraph(f"Late: {late_count}").bold = True

        # Add individual attendance records
        doc.add_paragraph("\nDetailed Attendance Records:")
        for index, row in student_attendance.iterrows():
            date = row['check_in_datetime']
            status = row['status']
            course_name = row['course_name']
            doc.add_paragraph(f"{date} - {course_name}: {status}")

        # Add footer
        footer = doc.sections[0].footer
        footer_paragraph = footer.paragraphs[0]
        footer_paragraph.text = "Attendance Report - Generated by Attendance Tracker"
        footer_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Save the document
        doc_name = f"Attendance_Report_{student_id}.docx"
        doc.save(doc_name)

        print(f"Attendance report for {first_name} {last_name} (ID: {student_id}) has been saved as {doc_name}.")
    else:
        print("Report not exported.")

# Function to generate a report by date
def report_by_date():
    # Load the 'Courses' sheet to dynamically pull the list of courses
    courses_df = pd.read_excel(file_path, sheet_name='Courses')
    course_names = courses_df['course_name'].tolist()

    # Ask the user to select a course from the list
    print("Select a course:")
    for index, course_name in enumerate(course_names, 1):
        print(f"{index}. {course_name}")
    course_choice = int(input("Enter the number of the course: ")) - 1

    if course_choice < 0 or course_choice >= len(course_names):
        print("Invalid course selection.")
        return

    class_name = course_names[course_choice]
    date = input("Enter date (YYYY-MM-DD): ")

    # Load the 'Attendance' sheet
    attendance_df = pd.read_excel(file_path, sheet_name='Attendance')

    # Filter attendance data for the given class and date
    class_attendance = attendance_df[(attendance_df['course_name'] == class_name) & 
                                     (attendance_df['check_in_datetime'].str.startswith(date))]

    if class_attendance.empty:
        print(f"No attendance records found for {class_name} on {date}.")
        return

    # Count attendance statistics for the date
    present_count = class_attendance[class_attendance['status'] == 'Present'].shape[0]
    absent_count = class_attendance[class_attendance['status'] == 'Absent'].shape[0]
    late_count = class_attendance[class_attendance['status'] == 'Late'].shape[0]

    # Display a nicely formatted report
    print(f"Attendance Report for {class_name} on {date}")
    print(f"Present: {present_count}")
    print(f"Absent: {absent_count}")
    print(f"Late: {late_count}")

def report_by_major():
    # Load the 'All students' and 'Attendance' sheets
    all_students_df = pd.read_excel(file_path, sheet_name='All students')
    attendance_df = pd.read_excel(file_path, sheet_name='Attendance')

    # Merge the attendance and student data by 'student_id'
    merged_df = pd.merge(attendance_df, all_students_df, on='student_id')

    # Get a list of majors
    majors = all_students_df['major'].unique().tolist()

    # Ask the user to select a major from the list
    print("Select a major from the list below:")
    for idx, major in enumerate(majors, 1):
        print(f"{idx}. {major}")
    major_choice = int(input("Enter the number of the major: ")) - 1

    if major_choice < 0 or major_choice >= len(majors):
        print("Invalid major selection.")
        return

    selected_major = majors[major_choice]
    print(f"Generating report for major: {selected_major}")

    # Filter by the selected major
    major_df = merged_df[merged_df['major'] == selected_major]

    # Count attendance statistics for the selected major
    attendance_report = major_df.groupby('status').size().reindex(['Present', 'Absent', 'Late'], fill_value=0)

    # Display the report in the application
    print(f"\nAttendance Report for Major: {selected_major}")
    print(f"Present: {attendance_report['Present']}")
    print(f"Absent: {attendance_report['Absent']}")
    print(f"Late: {attendance_report['Late']}")
    
    print("\nDetailed Attendance Records:")
    for index, row in major_df.iterrows():
        # Use '_x' columns from the Attendance DataFrame
        student_name = f"{row['first_name_x']} {row['last_name_x']}"
        date = row['check_in_datetime']
        status = row['status']
        course_name = row['course_name_x']
        print(f"{student_name} - {course_name} - {status} - {date}")

    # Ask if the user wants to export the report to a PDF
    export_choice = input("\nWould you like to export this report to a PDF? (y/n): ").lower()

    if export_choice == 'y':
        # Generate PDF Report
        pdf = FPDF()

        # Add a page
        pdf.add_page()

        # Set title
        pdf.set_font('Arial', 'B', 16)
        pdf.cell(200, 10, f"Attendance Report by Major: {selected_major}", ln=True, align='C')

        # Add attendance summary
        pdf.set_font('Arial', '', 12)
        pdf.cell(200, 10, 'Attendance Summary:', ln=True)
        pdf.cell(200, 10, f"Present: {attendance_report['Present']}", ln=True)
        pdf.cell(200, 10, f"Absent: {attendance_report['Absent']}", ln=True)
        pdf.cell(200, 10, f"Late: {attendance_report['Late']}", ln=True)

        # Add detailed attendance records
        pdf.cell(200, 10, '\nDetailed Attendance Records:', ln=True)
        for index, row in major_df.iterrows():
            # Use '_x' columns from the Attendance DataFrame
            student_name = f"{row['first_name_x']} {row['last_name_x']}"
            date = row['check_in_datetime']
            status = row['status']
            course_name = row['course_name_x']
            pdf.cell(200, 10, f"{student_name} - {course_name} - {status} - {date}", ln=True)

        # Save the PDF
        pdf_file_name = f"Attendance_Report_By_Major_{selected_major}.pdf"
        pdf.output(pdf_file_name)

        print(f"Attendance report for {selected_major} has been saved as {pdf_file_name}.")
    else:
        print("Report not exported.")


# Example usage of the report menu
def attendance_report_menu():
    while True:
        print("===== Attendance Report Menu =====")
        print("1. Report by Student")
        print("2. Report by Date")
        print("3. Report by Major")
        print("4. Exit")
        
        choice = input("Enter your choice: ")

        if choice == "1":
            report_by_student()
        elif choice == "2":
            report_by_date()
        elif choice == "3":
            report_by_major()
        elif choice == "4":
            print("Exiting report menu.")
            break
        else:
            print("Invalid choice. Please try again.")


# to delete a student and log when and fi deletion happend
def delete_student():
    print("===== Delete Student =====")
    student_id = input("Enter the student ID to delete: ")

    # Load the 'All students' sheet
    all_students_df = pd.read_excel(file_path, sheet_name='All students')

    # Check if the student ID exists in the DataFrame
    if student_id in all_students_df['student_id'].astype(str).values:
        # Mark the student as deleted and add a deletion date
        student_index = all_students_df[all_students_df['student_id'].astype(str) == student_id].index[0]
        all_students_df.at[student_index, 'is_deleted'] = True
        all_students_df.at[student_index, 'deletion_date'] = datetime.now().strftime("%Y-%m-%d %H:%M:%S")

        print(f"Student with ID {student_id} has been marked as deleted.")
        
        # Save the updated DataFrame back to the Excel file
        with pd.ExcelWriter(file_path, mode='a', engine='openpyxl', if_sheet_exists='replace') as writer:
            all_students_df.to_excel(writer, sheet_name='All students', index=False)

        print("The updated roster has been saved.")
    else:
        print(f"No student found with ID {student_id}.")

#updates the student information and ask for confirmation befor making changes.
def update_student_info():
    print("===== Update Student Information =====")
    student_id = input("Enter the student ID to update (or type 'cancel' to exit): ")
    if student_id.lower() == 'cancel':
        print("Operation canceled.")
        return

    # Load the 'All students' and 'Courses' sheets
    all_students_df = pd.read_excel(file_path, sheet_name='All students')
    courses_df = pd.read_excel(file_path, sheet_name='Courses')

    # Check if the student ID exists in the DataFrame
    if student_id in all_students_df['student_id'].astype(str).values:
        # Get the index of the student
        student_index = all_students_df[all_students_df['student_id'].astype(str) == student_id].index[0]

        # Display current information for the student
        print(f"Current Information for Student ID {student_id}:")
        print(all_students_df.loc[student_index])

        # Ask the user which field they want to update
        print("\nWhat would you like to update?")
        print("1. First Name")
        print("2. Last Name")
        print("3. Email")
        print("4. Major")
        print("5. Course Name")
        print("6. Cancel")
        choice = input("Enter the number of the field you want to update: ")
        
        # Handle cancel option immediately
        if choice == "6":
            print("Operation canceled.")
            return

        # Update first name
        if choice == "1":
            new_first_name = input("Enter new first name (or type 'cancel' to exit): ")
            if new_first_name.lower() == 'cancel':
                print("Change canceled.")
                return
            confirmation = input(f"Confirm change: Update First Name to '{new_first_name}'? (y/n): ").lower()
            if confirmation == 'y':
                all_students_df.at[student_index, 'first_name'] = new_first_name
                print(f"First name updated to {new_first_name}.")
            else:
                print("Change canceled.")

        # Update last name
        elif choice == "2":
            new_last_name = input("Enter new last name (or type 'cancel' to exit): ")
            if new_last_name.lower() == 'cancel':
                print("Change canceled.")
                return
            confirmation = input(f"Confirm change: Update Last Name to '{new_last_name}'? (y/n): ").lower()
            if confirmation == 'y':
                all_students_df.at[student_index, 'last_name'] = new_last_name
                print(f"Last name updated to {new_last_name}.")
            else:
                print("Change canceled.")

        # Update email
        elif choice == "3":
            new_email = input("Enter new email (or type 'cancel' to exit): ")
            if new_email.lower() == 'cancel':
                print("Change canceled.")
                return
            confirmation = input(f"Confirm change: Update Email to '{new_email}'? (y/n): ").lower()
            if confirmation == 'y':
                all_students_df.at[student_index, 'email'] = new_email
                print(f"Email updated to {new_email}.")
            else:
                print("Change canceled.")

        # Update major with a list of options
        elif choice == "4":
            majors = all_students_df['major'].unique().tolist()
            print("\nSelect a new major from the list (or type 'cancel' to exit):")
            for idx, major in enumerate(majors, 1):
                print(f"{idx}. {major}")
            major_choice = input("Enter the number of the new major: ")
            if major_choice.lower() == 'cancel':
                print("Change canceled.")
                return
            try:
                major_choice = int(major_choice) - 1
                if major_choice < 0 or major_choice >= len(majors):
                    print("Invalid major selection.")
                    return
                new_major = majors[major_choice]
                confirmation = input(f"Confirm change: Update Major to '{new_major}'? (y/n): ").lower()
                if confirmation == 'y':
                    all_students_df.at[student_index, 'major'] = new_major
                    print(f"Major updated to {new_major}.")
                else:
                    print("Change canceled.")
            except ValueError:
                print("Invalid input. Please enter a valid number.")

        # Update course name with a list of options
        elif choice == "5":
            courses = courses_df['course_name'].tolist()
            print("\nSelect a new course from the list (or type 'cancel' to exit):")
            for idx, course in enumerate(courses, 1):
                print(f"{idx}. {course}")
            course_choice = input("Enter the number of the new course: ")
            if course_choice.lower() == 'cancel':
                print("Change canceled.")
                return
            try:
                course_choice = int(course_choice) - 1
                if course_choice < 0 or course_choice >= len(courses):
                    print("Invalid course selection.")
                    return
                new_course_name = courses[course_choice]
                confirmation = input(f"Confirm change: Update Course Name to '{new_course_name}'? (y/n): ").lower()
                if confirmation == 'y':
                    all_students_df.at[student_index, 'course_name'] = new_course_name
                    print(f"Course name updated to {new_course_name}.")
                else:
                    print("Change canceled.")
            except ValueError:
                print("Invalid input. Please enter a valid number.")

        else:
            print("Invalid choice.")

        # Save the updated DataFrame back to the Excel file
        with pd.ExcelWriter(file_path, mode='a', engine='openpyxl', if_sheet_exists='replace') as writer:
            all_students_df.to_excel(writer, sheet_name='All students', index=False)

        print("The student's information has been updated and saved.")
    else:
        print(f"No student found with ID {student_id}.")

# Function to generate a student roster report for a class
def roster_report():
    print("===== Student Roster Report =====")
    
    # Load the 'All students' and 'Courses' sheets
    all_students_df = pd.read_excel(file_path, sheet_name='All students')
    courses_df = pd.read_excel(file_path, sheet_name='Courses')

    # Display the list of active courses for the user to select
    print("Select a course from the list below:")
    active_courses = courses_df['course_name'].tolist()

    for index, course in enumerate(active_courses, 1):
        print(f"{index}. {course}")

    try:
        course_choice = int(input("Enter the number of the course: ")) - 1

        if course_choice < 0 or course_choice >= len(active_courses):
            print("Invalid choice.")
            return
    except ValueError:
        print("Invalid input. Please enter a valid number.")
        return

    course_name = active_courses[course_choice].strip()  # Ensure there's no leading/trailing whitespace
    print(f"Generating roster for: {course_name}")

    # Ask the user if they want to include deleted students
    include_deleted = input("Include deleted students? (y/n): ").lower() == 'y'

    # Filter the students based on the course name
    roster_df = all_students_df[all_students_df['course_name'].str.strip() == course_name]

    # Exclude deleted students if the user chooses not to include them
    if not include_deleted:
        roster_df = roster_df[~(roster_df['is_deleted'] == True)]  # Exclude deleted students

    # Check if the roster is empty
    if roster_df.empty:
        print(f"No students found for the course: {course_name}")
        return

    # Display the roster
    print(f"\nStudent Roster for {course_name}:")
    print(roster_df[['student_id', 'first_name', 'last_name', 'email', 'major']])

    # If including deleted students, show the deletion date
    if include_deleted:
        deleted_students_df = all_students_df[(all_students_df['course_name'].str.strip() == course_name) & (all_students_df['is_deleted'] == True)]
        if not deleted_students_df.empty:
            print("\nDeleted Students:")
            print(deleted_students_df[['student_id', 'first_name', 'last_name', 'deletion_date']])
        else:
            print("\nNo deleted students for this class.")


# Function to display application information
def application_info():
    creator_name = "Stephon Dean"  # Replace with your name
    current_year = datetime.now().year
    copyright_symbol = "©"
    
    print(f"\n===== Application Information =====")
    print(f"Creator: {creator_name}")
    print(f"{copyright_symbol} {current_year} All Rights Reserved.")
    print("===================================")


# Function to load data from all relevant sheets in the Excel file
def load_all_data():
    # Load all sheets into a dictionary of DataFrames
    excel_data = pd.read_excel(file_path, sheet_name=None)
    
    # Convert each sheet into a list of dictionaries
    structured_data = {}
    for sheet_name, df in excel_data.items():
        structured_data[sheet_name] = df.to_dict(orient='records')
    
    return structured_data

# Initialize the OpenAI client for NVIDIA Llama API
##client = OpenAI(
    base_url="https://integrate.api.nvidia.com/v1",  # NVIDIA API base URL
    api_key="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"  # Replace with your actual API key
#)

# Function to filter the data based on the user's question
def filter_course_data(question, all_data):
    # Extract relevant information from the question
    days_of_interest = ["monday", "tuesday", "wednesday", "thursday", "friday"]
    relevant_days = [day for day in days_of_interest if day in question.lower()]

    # If the question mentions specific days, filter courses that include those days
    if relevant_days:
        course_data = all_data.get('Courses', [])
        filtered_courses = [
            course for course in course_data
            if any(day.lower() in course['meeting_date'].lower() for day in relevant_days)
        ]
        return filtered_courses

    # Default to returning all courses if no specific days are mentioned
    return all_data.get('Courses', [])

# Function to query the NVIDIA Llama API with the user's question and filtered dataset
def query_llama_api(question, all_data):
    # Filter course data based on the question to reduce input size
    filtered_data = filter_course_data(question, all_data)
    if not filtered_data:
        return "No relevant course data found for your query."

    # Prepare a more concise prompt with the filtered data
    prompt = f"Using the following filtered course data: {filtered_data}, answer the question: {question}"

    # Send the prompt to the NVIDIA Llama API
    completion = client.chat.completions.create(
        model="meta/llama-3.1-405b-instruct",  # Use the specific Llama model version
        messages=[{"role": "user", "content": prompt}],
        temperature=0.2,  # Adjust the temperature as needed
        top_p=0.7,
        max_tokens=1024,
        stream=True  # Enable streaming
    )

    # Handle the streaming response and capture the output
    answer = ""
    for chunk in completion:
        if chunk.choices[0].delta.content is not None:
            answer += chunk.choices[0].delta.content

    return answer.strip()

# Function to handle the AI course inquiry process
def ai_course_inquiry():
    print("===== AI Course Inquiry =====")
    
    # Load all data from the Excel file
    all_data = load_all_data()

    while True:
        # Ask the user for their question
        question = input("\nAsk a question about the courses or students (or type 'back' to return to the main menu): ")

        # Check if the user wants to return to the main menu
        if question.lower() == 'back':
            print("Returning to the main menu...")
            break

        # Query the NVIDIA Llama API for an answer based on the filtered dataset
        answer = query_llama_api(question, all_data)

        # Display the AI's answer
        print(f"AI Answer: {answer}")

def generate_charts_menu():
    while True:
        print("\n===== Generate Charts and Graphs =====")
        print("1. Attendance Summary by Course (Bar Chart)")
        print("2. Attendance Pie Chart (Pie Chart)")
        print("3. Attendance by Major (Stacked Bar Chart)")
        print("4. Back to Main Menu")
        chart_choice = input("Enter your choice: ")

        if chart_choice == "1":
            generate_attendance_by_course_chart()
        elif chart_choice == "2":
            generate_attendance_pie_chart()
        elif chart_choice == "3":
            generate_attendance_by_major_chart()
        elif chart_choice == "4":
            print("Returning to the main menu...")
            break
        else:
            print("Invalid choice. Please select again.")


def generate_attendance_by_course_chart():
    # Load the attendance data
    attendance_df = pd.read_excel(file_path, sheet_name='Attendance')

    # Group by course and status to get attendance counts
    course_attendance = attendance_df.groupby(['course_name', 'status']).size().unstack(fill_value=0)

    # Plot the data
    course_attendance.plot(kind='bar', stacked=True, color=['green', 'red', 'orange'])

    # Add labels and title
    plt.title('Attendance Summary by Course')
    plt.xlabel('Course Name')
    plt.ylabel('Number of Students')
    plt.legend(title="Attendance Status")

    # Show the plot
    plt.show()

def generate_attendance_pie_chart():
    # Load the attendance data
    attendance_df = pd.read_excel(file_path, sheet_name='Attendance')

    # Convert 'check_in_datetime' to datetime and extract just the date part
    attendance_df['check_in_datetime'] = pd.to_datetime(attendance_df['check_in_datetime'])
    attendance_df['date'] = attendance_df['check_in_datetime'].dt.date

    # Ask the user for a specific date
    date_input = input("Enter the date for the attendance report (YYYY-MM-DD): ")
    try:
        specific_date = pd.to_datetime(date_input).date()
    except ValueError:
        print("Invalid date format. Please enter the date in the format YYYY-MM-DD.")
        return

    # Filter the data for the specific date
    specific_date_data = attendance_df[attendance_df['date'] == specific_date]

    # Check if there's any attendance data for the selected date
    if specific_date_data.empty:
        print(f"No attendance data available for {date_input}.")
        return

    # Get attendance counts for each status (Present, Absent, Late)
    attendance_counts = specific_date_data['status'].value_counts()

    # Plot the pie chart
    attendance_counts.plot(kind='pie', autopct='%1.1f%%', startangle=90, colors=['green', 'red', 'orange'])

    # Add title and adjust layout
    plt.title(f'Attendance Distribution for {specific_date}')
    plt.ylabel('')  # Hide the y-label

    plt.tight_layout()
    plt.show()


def generate_attendance_by_major_chart():
    # Load the attendance and student data
    attendance_df = pd.read_excel(file_path, sheet_name='Attendance')
    all_students_df = pd.read_excel(file_path, sheet_name='All students')

    # Merge the attendance data with the student data to get major information
    merged_df = pd.merge(attendance_df, all_students_df, on='student_id')

    # Group by major and status to get attendance counts
    major_attendance = merged_df.groupby(['major', 'status']).size().unstack(fill_value=0)

    # Plot the data
    major_attendance.plot(kind='bar', stacked=True, color=['green', 'red', 'orange'])

    # Add labels and title
    plt.title('Attendance by Major')
    plt.xlabel('Major')
    plt.ylabel('Number of Students')
    plt.legend(title="Attendance Status")

    # Show the plot
    plt.show()





def logged_in_menu(username):
    while True:
        print(f"===== {username}'s Menu =====")
        print("1. Take attendance")
        print("2. Add course")
        print("3. Add student")
        print("4. Delete student")
        print("5. Update student information")
        print("6. Run reports")
        print("7. Generate student roster report")
        print("8. Application Information")
        print("9. AI Course Inquiry")
        print("10. Generate Charts and Graphs")  # New option for charts and graphs
        print("11. Log out")
        print("12. Exit the application")
        choice = input("Enter your choice: ")

        if choice == "1":
            take_attendance()
        elif choice == "2":
            add_course()
        elif choice == "3":
            add_student()
        elif choice == "4":
            delete_student()
        elif choice == "5":
            update_student_info()
        elif choice == "6":
            attendance_report_menu()
        elif choice == "7":
            roster_report()
        elif choice == "8":
            application_info()
        elif choice == "9":
            ai_course_inquiry()
        elif choice == "10":  # Generate Charts and Graphs
            generate_charts_menu()
        elif choice == "11":
            print(f"Logging out {username}...")
            break  # Log out and return to the login menu
        elif choice == "12":
            print("Exiting the application...")
            exit()  # Exit the application entirely
        else:
            print("Invalid choice. Please select again.")


# Main function to run the application
if __name__ == "__main__":
    login()

   
